"""Сборка DOCX для отчёта по сохранённому дайджесту (python-docx)."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from digest.models import DigestResponse


def saved_digest_to_docx_bytes(
    title: str,
    created_at: str,
    response: DigestResponse,
    topic_queries: list[str] | None,
) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(11)

    doc.add_heading(title, 0)
    doc.add_paragraph(f"Создано (UTC): {created_at}")

    if topic_queries:
        doc.add_paragraph("Темы запроса: " + " · ".join(topic_queries))

    meta = response.meta
    if meta:
        doc.add_heading("Мета", level=1)
        lines = [
            f"Режим: {meta.digest_mode}",
            f"В LLM: {meta.used_for_llm}",
            f"Секунд: {meta.elapsed_seconds:.1f}" if meta.elapsed_seconds else None,
        ]
        doc.add_paragraph(" · ".join(x for x in lines if x))

    doc.add_heading("Дайджест (RU)", level=1)
    for block in (response.digest_ru or "—").split("\n"):
        doc.add_paragraph(block.strip() or " ")

    doc.add_heading("Digest (EN)", level=1)
    for block in (response.digest_en or "—").split("\n"):
        doc.add_paragraph(block.strip() or " ")

    doc.add_heading("Публикации", level=1)
    for i, p in enumerate(response.publications_used, 1):
        head = f"{i}. {p.title}"
        if p.year is not None:
            head += f" ({p.year})"
        para = doc.add_paragraph()
        run = para.add_run(head)
        run.bold = True
        bits: list[str] = []
        if p.doi:
            bits.append(f"DOI: {p.doi}")
        if p.url:
            bits.append(p.url)
        if p.source:
            bits.append(p.source)
        if bits:
            doc.add_paragraph(" · ".join(bits))
        if (p.abstract or "").strip():
            doc.add_paragraph((p.abstract or "").strip()[:4000])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
